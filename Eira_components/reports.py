from .database import DB_PATH
import sqlite3
import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
from tkinter import filedialog

def generate_word_report(patient_id, user_id):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    # Get patient info
    cursor.execute("SELECT * FROM patients WHERE id = ?", (patient_id,))
    patient = dict(cursor.fetchone()) if cursor.fetchone() else None

    if not patient:
        conn.close()
        return {"success": False, "message": "Patient not found"}

    # Get therapist info
    cursor.execute("SELECT * FROM users WHERE id = ?", (user_id,))
    therapist = dict(cursor.fetchone()) if cursor.fetchone() else None

    # Get SOAP notes
    cursor.execute(
        """SELECT * FROM soap_notes 
        WHERE patient_id = ? 
        ORDER BY date DESC""",
        (patient_id,)
    )
    soap_notes = [dict(row) for row in cursor.fetchall()]

    conn.close()

    if not patient or not therapist:
        return {"success": False, "message": "Could not find patient or therapist information"}

    # Create a new Document
    doc = Document()

    # Set margins for the document
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run('Physical Therapy Progress Report')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add date
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f"Date Generated: {datetime.datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(10)

    doc.add_paragraph()

    # Therapist Information
    doc.add_heading('Therapist Information', level=2)
    therapist_table = doc.add_table(rows=2, cols=2)
    therapist_table.style = 'Table Grid'

    # Add therapist details
    therapist_table.cell(0, 0).text = "Name:"
    therapist_table.cell(0, 1).text = f"{therapist['first_name']} {therapist['last_name']}"
    therapist_table.cell(1, 0).text = "Email:"
    therapist_table.cell(1, 1).text = therapist['email']

    doc.add_paragraph()

    # Patient Information
    doc.add_heading('Patient Information', level=2)
    patient_table = doc.add_table(rows=8, cols=2)
    patient_table.style = 'Table Grid'

    # Calculate age
    dob = datetime.datetime.strptime(patient['date_of_birth'], '%Y-%m-%d').date()
    today = datetime.date.today()
    age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))

    # Add patient details
    patient_table.cell(0, 0).text = "Name:"
    patient_table.cell(0, 1).text = f"{patient['first_name']} {patient['last_name']}"
    patient_table.cell(1, 0).text = "DOB:"
    patient_table.cell(1, 1).text = datetime.datetime.strptime(patient['date_of_birth'], '%Y-%m-%d').strftime('%B %d, %Y')
    patient_table.cell(2, 0).text = "Age:"
    patient_table.cell(2, 1).text = str(age)
    patient_table.cell(3, 0).text = "Gender:"
    patient_table.cell(3, 1).text = patient['gender'].capitalize() if patient['gender'] else "Not specified"
    patient_table.cell(4, 0).text = "Diagnosis:"
    patient_table.cell(4, 1).text = patient['diagnosis']
    patient_table.cell(5, 0).text = "Medical Aid:"
    patient_table.cell(5, 1).text = patient['medical_aid_name'] or "Not provided"
    patient_table.cell(6, 0).text = "Medical Aid #:"
    patient_table.cell(6, 1).text = patient['medical_aid_number'] or "Not provided"

    doc.add_paragraph()

    # Medical History
    doc.add_heading('Medical History', level=2)
    doc.add_paragraph(patient['medical_history'] if patient['medical_history'] else "No medical history recorded.")

    doc.add_paragraph()

    # Treatment Summary & Progress
    doc.add_heading('Treatment Summary & Progress', level=2)

    if not soap_notes:
        doc.add_paragraph("No treatment notes available for this patient.")
    else:
        # Get first and latest note for comparison
        first_note = soap_notes[-1] if soap_notes else None
        latest_note = soap_notes[0] if soap_notes else None

        # Initial assessment
        if first_note:
            doc.add_heading('Initial Assessment', level=3)
            first_date = datetime.datetime.strptime(first_note['date'], '%Y-%m-%d %H:%M:%S').strftime('%B %d, %Y')
            doc.add_paragraph(f"Date: {first_date}")

            initial_table = doc.add_table(rows=4, cols=2)
            initial_table.style = 'Table Grid'

            initial_table.cell(0, 0).text = "Subjective:"
            initial_table.cell(0, 1).text = first_note['subjective']
            initial_table.cell(1, 0).text = "Objective:"
            initial_table.cell(1, 1).text = first_note['objective']
            initial_table.cell(2, 0).text = "Action:"
            initial_table.cell(2, 1).text = first_note['action']
            initial_table.cell(3, 0).text = "Plan:"
            initial_table.cell(3, 1).text = first_note['plan']

            doc.add_paragraph()

        # Latest assessment (if different from first)
        if latest_note and latest_note['id'] != first_note['id']:
            doc.add_heading('Most Recent Assessment', level=3)
            latest_date = datetime.datetime.strptime(latest_note['date'], '%Y-%m-%d %H:%M:%S').strftime('%B %d, %Y')
            doc.add_paragraph(f"Date: {latest_date}")

            latest_table = doc.add_table(rows=4, cols=2)
            latest_table.style = 'Table Grid'

            latest_table.cell(0, 0).text = "Subjective:"
            latest_table.cell(0, 1).text = latest_note['subjective']
            latest_table.cell(1, 0).text = "Objective:"
            latest_table.cell(1, 1).text = latest_note['objective']
            latest_table.cell(2, 0).text = "Action:"
            latest_table.cell(2, 1).text = latest_note['action']
            latest_table.cell(3, 0).text = "Plan:"
            latest_table.cell(3, 1).text = latest_note['plan']

            doc.add_paragraph()

        # Treatment progress
        doc.add_heading('Treatment Progress', level=3)
        if len(soap_notes) > 1:
            progress_paragraph = doc.add_paragraph()
            progress_run = progress_paragraph.add_run(f"Total sessions: {len(soap_notes)}")
            progress_run.bold = True

            first_date = datetime.datetime.strptime(soap_notes[-1]['date'], '%Y-%m-%d %H:%M:%S').strftime('%B %d, %Y')
            last_date = datetime.datetime.strptime(soap_notes[0]['date'], '%Y-%m-%d %H:%M:%S').strftime('%B %d, %Y')
            doc.add_paragraph(f"Treatment period: {first_date} to {last_date}")

            # Extract goals progress from the most recent note if available
            if latest_note and latest_note['goals_progress']:
                goals_para = doc.add_paragraph()
                goals_run = goals_para.add_run("Goals Progress:")
                goals_run.bold = True
                doc.add_paragraph(latest_note['goals_progress'])
        else:
            doc.add_paragraph("Insufficient data to determine progress. Only one session recorded.")

    # Create a buffer to save the file
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Save the file to a location chosen by the user
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word documents", "*.docx")],
        initialfile=f"Progress_Report_{patient['last_name']}_{patient['first_name']}_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
    )

    if file_path:
        with open(file_path, 'wb') as f:
            f.write(buffer.getvalue())
        return {"success": True, "message": f"Report saved to {file_path}"}
    else:
        return {"success": False, "message": "Report generation cancelled"}

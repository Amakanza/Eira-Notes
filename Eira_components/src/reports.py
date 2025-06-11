import datetime
import io
import sqlite3
from tkinter import filedialog
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .database import DB_PATH


def generate_word_report(patient_id, user_id):
    """Create and save a Word progress report for a patient."""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    cursor.execute("SELECT * FROM patients WHERE id = ?", (patient_id,))
    patient_row = cursor.fetchone()
    patient = dict(patient_row) if patient_row else None
    if not patient:
        conn.close()
        return {"success": False, "message": "Patient not found"}

    cursor.execute("SELECT * FROM users WHERE id = ?", (user_id,))
    therapist_row = cursor.fetchone()
    therapist = dict(therapist_row) if therapist_row else None

    cursor.execute(
        "SELECT * FROM soap_notes WHERE patient_id = ? ORDER BY date DESC",
        (patient_id,),
    )
    soap_notes = [dict(row) for row in cursor.fetchall()]
    conn.close()

    if not therapist:
        return {"success": False, "message": "Therapist not found"}

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    title = doc.add_paragraph()
    run = title.add_run("Physical Therapy Progress Report")
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(
        f"Date Generated: {datetime.datetime.now().strftime('%B %d, %Y')}"
    )
    date_run.font.size = Pt(10)
    doc.add_paragraph()

    doc.add_heading("Therapist Information", level=2)
    t_table = doc.add_table(rows=2, cols=2)
    t_table.style = "Table Grid"
    t_table.cell(0, 0).text = "Name:"
    t_table.cell(0, 1).text = f"{therapist['first_name']} {therapist['last_name']}"
    t_table.cell(1, 0).text = "Email:"
    t_table.cell(1, 1).text = therapist['email']
    doc.add_paragraph()

    doc.add_heading("Patient Information", level=2)
    p_table = doc.add_table(rows=7, cols=2)
    p_table.style = "Table Grid"
    dob = datetime.datetime.strptime(patient["date_of_birth"], "%Y-%m-%d").date()
    today = datetime.date.today()
    age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
    p_table.cell(0, 0).text = "Name:"
    p_table.cell(0, 1).text = f"{patient['first_name']} {patient['last_name']}"
    p_table.cell(1, 0).text = "DOB:"
    p_table.cell(1, 1).text = dob.strftime("%B %d, %Y")
    p_table.cell(2, 0).text = "Age:"
    p_table.cell(2, 1).text = str(age)
    p_table.cell(3, 0).text = "Gender:"
    p_table.cell(3, 1).text = patient.get("gender", "").capitalize() or "Not specified"
    p_table.cell(4, 0).text = "Diagnosis:"
    p_table.cell(4, 1).text = patient['diagnosis']
    p_table.cell(5, 0).text = "Medical Aid:"
    p_table.cell(5, 1).text = patient['medical_aid_name'] or "Not provided"
    p_table.cell(6, 0).text = "Medical Aid #:"
    p_table.cell(6, 1).text = patient['medical_aid_number'] or "Not provided"
    doc.add_paragraph()

    doc.add_heading("Medical History", level=2)
    doc.add_paragraph(patient.get("medical_history") or "No medical history recorded.")
    doc.add_paragraph()

    doc.add_heading("Treatment Summary & Progress", level=2)

    if not soap_notes:
        doc.add_paragraph("No treatment notes available for this patient.")
    else:
        first_note = soap_notes[-1] if soap_notes else None
        latest_note = soap_notes[0] if soap_notes else None
        if first_note:
            doc.add_heading("Initial Assessment", level=3)
            fd = datetime.datetime.strptime(first_note['date'], '%Y-%m-%d %H:%M:%S').strftime('%B %d, %Y')
            doc.add_paragraph(f"Date: {fd}")
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'
            table.cell(0, 0).text = "Subjective:"
            table.cell(0, 1).text = first_note['subjective']
            table.cell(1, 0).text = "Objective:"
            table.cell(1, 1).text = first_note['objective']
            table.cell(2, 0).text = "Action:"
            table.cell(2, 1).text = first_note['action']
            table.cell(3, 0).text = "Plan:"
            table.cell(3, 1).text = first_note['plan']
            doc.add_paragraph()
        if latest_note and latest_note['id'] != first_note['id']:
            doc.add_heading("Most Recent Assessment", level=3)
            ld = datetime.datetime.strptime(latest_note['date'], '%Y-%m-%d %H:%M:%S').strftime('%B %d, %Y')
            doc.add_paragraph(f"Date: {ld}")
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'
            table.cell(0, 0).text = "Subjective:"
            table.cell(0, 1).text = latest_note['subjective']
            table.cell(1, 0).text = "Objective:"
            table.cell(1, 1).text = latest_note['objective']
            table.cell(2, 0).text = "Action:"
            table.cell(2, 1).text = latest_note['action']
            table.cell(3, 0).text = "Plan:"
            table.cell(3, 1).text = latest_note['plan']
            doc.add_paragraph()
        doc.add_heading("Treatment Progress", level=3)
        if len(soap_notes) > 1:
            ppara = doc.add_paragraph()
            run = ppara.add_run(f"Total sessions: {len(soap_notes)}")
            run.bold = True
            first_date = datetime.datetime.strptime(soap_notes[-1]['date'], '%Y-%m-%d %H:%M:%S').strftime('%B %d, %Y')
            last_date = datetime.datetime.strptime(soap_notes[0]['date'], '%Y-%m-%d %H:%M:%S').strftime('%B %d, %Y')
            doc.add_paragraph(f"Treatment period: {first_date} to {last_date}")
            if latest_note and latest_note.get('goals_progress'):
                gpara = doc.add_paragraph()
                run = gpara.add_run("Goals Progress:")
                run.bold = True
                doc.add_paragraph(latest_note['goals_progress'])
        else:
            doc.add_paragraph("Insufficient data to determine progress. Only one session recorded.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word documents", "*.docx")],
        initialfile=f"Progress_Report_{patient['last_name']}_{patient['first_name']}_{datetime.datetime.now().strftime('%Y%m%d')}.docx",
    )
    if file_path:
        with open(file_path, "wb") as f:
            f.write(buffer.getvalue())
        return {"success": True, "message": f"Report saved to {file_path}"}
    return {"success": False, "message": "Report generation cancelled"}

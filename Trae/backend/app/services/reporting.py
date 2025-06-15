import os
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

from docx import Document
from jinja2 import Template
from weasyprint import HTML

from app.db.models.clinical_note import ClinicalNote
from app.db.models.patient import Patient
from app.db.models.user import User
from app.services.storage import storage_service


class ReportingService:
    """Service for generating clinical reports."""
    
    def __init__(self):
        """Initialize the reporting service."""
        self.templates_dir = Path("app/templates/reports")
        os.makedirs(self.templates_dir, exist_ok=True)
    
    async def generate_docx_report(
        self,
        template_name: str,
        context: Dict[str, Any],
        output_filename: Optional[str] = None,
    ) -> str:
        """Generate a DOCX report from a template and context."""
        # Load the template document
        template_path = self.templates_dir / f"{template_name}.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"Template {template_name}.docx not found")
        
        doc = Document(template_path)
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in context.items():
                if f"{{{{ {key} }}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        f"{{{{ {key} }}}}", str(value)
                    )
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in context.items():
                            if f"{{{{ {key} }}}}" in paragraph.text:
                                paragraph.text = paragraph.text.replace(
                                    f"{{{{ {key} }}}}", str(value)
                                )
        
        # Save the document to a temporary file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        
        doc.save(tmp_path)
        
        # Upload to S3
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            output_filename = f"{template_name}_{timestamp}.docx"
        
        with open(tmp_path, "rb") as f:
            upload_file = UploadFile(
                filename=output_filename,
                file=f,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            file_path, _ = await storage_service.upload_file(
                upload_file, folder="reports/docx"
            )
        
        # Clean up the temporary file
        os.unlink(tmp_path)
        
        return file_path
    
    async def generate_pdf_report(
        self,
        template_name: str,
        context: Dict[str, Any],
        output_filename: Optional[str] = None,
    ) -> str:
        """Generate a PDF report from an HTML template and context."""
        # Load the template
        template_path = self.templates_dir / f"{template_name}.html"
        if not template_path.exists():
            raise FileNotFoundError(f"Template {template_name}.html not found")
        
        with open(template_path, "r") as f:
            template_content = f.read()
        
        # Render the template with context
        template = Template(template_content)
        html_content = template.render(**context)
        
        # Generate PDF
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp_path = tmp.name
        
        HTML(string=html_content).write_pdf(tmp_path)
        
        # Upload to S3
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            output_filename = f"{template_name}_{timestamp}.pdf"
        
        with open(tmp_path, "rb") as f:
            upload_file = UploadFile(
                filename=output_filename,
                file=f,
                content_type="application/pdf",
            )
            file_path, _ = await storage_service.upload_file(
                upload_file, folder="reports/pdf"
            )
        
        # Clean up the temporary file
        os.unlink(tmp_path)
        
        return file_path
    
    async def generate_clinical_note_report(
        self, clinical_note: ClinicalNote, format_type: str = "pdf"
    ) -> str:
        """Generate a report for a clinical note."""
        # Prepare context
        patient = clinical_note.patient
        clinician = clinical_note.created_by
        
        context = {
            "note_title": clinical_note.title,
            "note_content": clinical_note.content,
            "note_type": clinical_note.note_type,
            "note_date": clinical_note.created_at.strftime("%Y-%m-%d %H:%M"),
            "patient_name": patient.full_name,
            "patient_dob": patient.date_of_birth.strftime("%Y-%m-%d"),
            "patient_age": patient.age,
            "patient_gender": patient.gender,
            "patient_mrn": patient.medical_record_number,
            "clinician_name": clinician.full_name,
            "generated_date": datetime.now().strftime("%Y-%m-%d %H:%M"),
        }
        
        # Generate report based on format
        if format_type.lower() == "pdf":
            return await self.generate_pdf_report(
                "clinical_note", context, f"clinical_note_{clinical_note.id}.pdf"
            )
        elif format_type.lower() == "docx":
            return await self.generate_docx_report(
                "clinical_note", context, f"clinical_note_{clinical_note.id}.docx"
            )
        else:
            raise ValueError(f"Unsupported format type: {format_type}")


# Create a singleton instance
reporting_service = ReportingService()


# Import here to avoid circular imports
from fastapi import UploadFile